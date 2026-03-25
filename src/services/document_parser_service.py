import os
import PyPDF2
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentParserService:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        extracted_text = ""
        # 1. PyMuPDFでの抽出
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text = page.get_text()
                if text:
                    extracted_text += text + "\n"
            doc.close()
            if extracted_text.strip():
                return extracted_text.strip()
        except Exception as e:
            logger.warning(f"PyMuPDFでのテキスト抽出に失敗しました: {e}")

        # 2. pdfplumberでの抽出（フォールバック）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
            if extracted_text.strip():
                return extracted_text.strip()
        except Exception as e:
            logger.warning(f"pdfplumberでの抽出に失敗しました: {e}")

        # 3. PyPDF2での抽出（最終フォールバック）
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
            return extracted_text.strip()
        except Exception as e:
            logger.error(f"PyPDF2でのテキスト抽出に失敗しました: {e}")
            raise Exception(f"PDFからのテキスト抽出に失敗しました: {e}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        full_text.append(" | ".join(row_data))
                        
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Word文書からのテキスト抽出に失敗しました: {e}")
            raise Exception(f"Word文書からのテキスト抽出に失敗しました: {e}")
